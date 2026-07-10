"""Парсеры для различных форматов документов."""

import os
from abc import ABC, abstractmethod
from typing import Optional


class ParseResult:
    """Результат парсинга документа."""

    def __init__(self, text: str, metadata: Optional[dict] = None):
        self.text = text
        self.metadata = metadata or {}


class DocumentParser(ABC):
    """Абстрактный парсер документов."""

    @abstractmethod
    def parse(self, file_path: str) -> ParseResult:
        """Парсинг документа, возврат текста и метаданных."""
        ...


class TxtParser(DocumentParser):
    """Парсер для TXT файлов."""

    def parse(self, file_path: str) -> ParseResult:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()
        return ParseResult(
            text=text,
            metadata={"source": file_path, "format": "txt"},
        )


class PdfParser(DocumentParser):
    """Парсер для PDF файлов (через PyMuPDF)."""

    def parse(self, file_path: str) -> ParseResult:
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise ImportError("PyMuPDF (fitz) is required for PDF parsing. Install: pip install pymupdf")

        doc = fitz.open(file_path)
        text_parts = []
        for page_num, page in enumerate(doc):
            text = page.get_text()
            text_parts.append(f"--- Page {page_num + 1} ---\n{text}")

        doc.close()
        return ParseResult(
            text="\n\n".join(text_parts),
            metadata={"source": file_path, "format": "pdf", "pages": len(doc)},
        )


class DocxParser(DocumentParser):
    """Парсер для DOCX файлов."""

    def parse(self, file_path: str) -> ParseResult:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("python-docx is required for DOCX parsing. Install: pip install python-docx")

        doc = DocxDocument(file_path)
        text_parts = []

        for para in doc.paragraphs:
            text_parts.append(para.text)

        # Также извлекаем таблицы
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                text_parts.append(row_text)

        return ParseResult(
            text="\n".join(text_parts),
            metadata={"source": file_path, "format": "docx"},
        )


class XlsxParser(DocumentParser):
    """Парсер для XLSX файлов (конвертация таблиц в Markdown)."""

    def parse(self, file_path: str) -> ParseResult:
        try:
            import openpyxl
        except ImportError:
            raise ImportError("openpyxl is required for XLSX parsing. Install: pip install openpyxl")

        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
        text_parts = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            text_parts.append(f"## Sheet: {sheet_name}")

            # Конвертируем таблицу в Markdown
            rows = []
            for row in ws.iter_rows(values_only=True):
                row_data = [str(cell) if cell is not None else "" for cell in row]
                rows.append(row_data)

            if rows:
                # Заголовок
                header = "| " + " | ".join(rows[0]) + " |"
                separator = "| " + " | ".join(["---"] * len(rows[0])) + " |"
                text_parts.append(header)
                text_parts.append(separator)

                # Данные
                for row in rows[1:]:
                    text_parts.append("| " + " | ".join(row) + " |")

            text_parts.append("")

        wb.close()
        return ParseResult(
            text="\n".join(text_parts),
            metadata={"source": file_path, "format": "xlsx", "sheets": wb.sheetnames},
        )


def get_parser(file_path: str) -> DocumentParser:
    """Фабрика: возвращает парсер в зависимости от расширения файла."""
    ext = os.path.splitext(file_path)[1].lower()

    parsers = {
        ".txt": TxtParser(),
        ".pdf": PdfParser(),
        ".docx": DocxParser(),
        ".doc": DocxParser(),
        ".xlsx": XlsxParser(),
        ".xls": XlsxParser(),
    }

    parser = parsers.get(ext)
    if not parser:
        raise ValueError(f"Unsupported file format: {ext}")

    return parser


def parse_document(file_path: str) -> ParseResult:
    """Парсинг документа по его пути."""
    parser = get_parser(file_path)
    return parser.parse(file_path)